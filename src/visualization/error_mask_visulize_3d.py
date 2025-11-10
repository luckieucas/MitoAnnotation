#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import re
import sys
import argparse
import numpy as np
import tifffile
from concurrent.futures import ThreadPoolExecutor
from PIL import Image, ImageDraw, ImageFont
import PIL
from typing import Union, Optional
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# VTK imports
import vtk
from vtkmodules.util import numpy_support

# Color map
LABEL_COLORMAP = [
    (255,   0,   0),
    (  0, 255,   0),
    (  0,   0, 255),
    (255, 255,   0),
    (255,   0, 255),
    (  0, 255, 255),
]

def relabel_data(data: np.ndarray):
    """Remap all non-zero labels to 1,2,3... and return new data"""
    unique = np.unique(data)
    unique = unique[unique != 0]
    mapping = {old: new for new, old in enumerate(unique, start=1)}
    out = np.zeros_like(data, dtype=np.uint8)
    for old, new in mapping.items():
        out[data == old] = new
    return out

def numpy_to_vtk_image(data: np.ndarray, spacing: tuple = (1.0, 1.0, 1.0)) -> vtk.vtkImageData:
    """Convert 3D numpy array to vtkImageData with specified spacing."""
    data = np.ascontiguousarray(data)
    depth, height, width = data.shape
    vtk_img = vtk.vtkImageData()
    vtk_img.SetDimensions(width, height, depth)
    vtk_img.SetSpacing(spacing[2], spacing[1], spacing[0])  # VTK uses (x, y, z) order
    vtk_array = numpy_support.numpy_to_vtk(
        num_array=data.ravel(order='C'),
        deep=True,
        array_type=vtk.VTK_UNSIGNED_CHAR
    )
    vtk_img.GetPointData().SetScalars(vtk_array)
    return vtk_img

def process_label(label: int, data: np.ndarray, smooth_iters: int, spacing: tuple = (1.0, 1.0, 1.0)) -> vtk.vtkActor:
    """Extract isosurface for a single label and return corresponding Actor."""
    if label == 0:
        return None
    # Create binary mask
    mask = np.where(data == label, label, 0).astype(np.uint8)
    vtk_img = numpy_to_vtk_image(mask, spacing)

    # Marching Cubes extract isosurface
    mc = vtk.vtkMarchingCubes()
    mc.SetInputData(vtk_img)
    mc.SetValue(0, label)
    mc.ComputeNormalsOn()
    mc.Update()

    # Optional smoothing
    if smooth_iters > 0:
        smoother = vtk.vtkSmoothPolyDataFilter()
        smoother.SetInputConnection(mc.GetOutputPort())
        smoother.SetNumberOfIterations(smooth_iters)
        smoother.SetRelaxationFactor(0.1)
        smoother.FeatureEdgeSmoothingOff()
        smoother.BoundarySmoothingOn()
        smoother.Update()
        poly = smoother.GetOutput()
    else:
        poly = mc.GetOutput()

    mapper = vtk.vtkPolyDataMapper()
    mapper.SetInputData(poly)
    mapper.ScalarVisibilityOff()

    actor = vtk.vtkActor()
    actor.SetMapper(mapper)
    r, g, b = LABEL_COLORMAP[label % len(LABEL_COLORMAP)]
    actor.GetProperty().SetColor(r/255.0, g/255.0, b/255.0)
    actor.GetProperty().SetOpacity(1.0)
    return actor

def render_single_view(data: np.ndarray, smooth_iters: int,
                       az: float, el: float,
                       size=(800,600), spacing: tuple = (1.0, 1.0, 1.0)) -> Image.Image:
    """
    Off-screen render a 3D mask from one viewpoint, return PIL.Image.
    az: horizontal rotation angle, el: elevation angle.
    spacing: voxel spacing in (z, y, x) order (depth, height, width).
    """
    # relabel data for better visualization
    data = relabel_data(data)
    # VTK renderer and window
    renderer = vtk.vtkRenderer()
    renderer.SetBackground(1,1,1)
    renwin = vtk.vtkRenderWindow()
    renwin.OffScreenRenderingOn()
    renwin.AddRenderer(renderer)
    renwin.SetSize(*size)

    # Add actors for all labels
    labels = np.unique(data)
    if len(labels) == 1:
        print("[WARN] Data contains only background labels, cannot render.")
        return Image.new('RGB', size, (255, 255, 255))
    with ThreadPoolExecutor() as exe:
        futures = {exe.submit(process_label, lab, data, smooth_iters, spacing): lab
                   for lab in labels if lab != 0}
        for f in futures:
            actor = f.result()
            if actor:
                renderer.AddActor(actor)

    if renderer.GetActors().GetNumberOfItems() == 0:
        raise RuntimeError("No renderable labels detected.")

    # Set camera viewpoint
    cam = renderer.GetActiveCamera()
    renderer.ResetCamera()
    cam.Azimuth(az)
    cam.Elevation(el)
    renderer.ResetCameraClippingRange()
    renwin.Render()

    # Capture image
    w2if = vtk.vtkWindowToImageFilter()
    w2if.SetInput(renwin)
    w2if.SetInputBufferTypeToRGB()
    w2if.ReadFrontBufferOff()
    w2if.Update()
    vtk_img = w2if.GetOutput()

    # Convert to numpy then PIL.Image
    width, height, _ = vtk_img.GetDimensions()
    arr = numpy_support.vtk_to_numpy(vtk_img.GetPointData().GetScalars())
    arr = arr.reshape(height, width, 3)
    return Image.fromarray(arr)

def combine_modalities(images, out_path: str):
    """Horizontally concatenate two PIL.Images and save to out_path."""
    w, h = images[0].size
    grid = Image.new('RGB', (w*2, h), (255,255,255))
    for i, img in enumerate(images):
        grid.paste(img, (i*w, 0))
    grid.save(out_path)
    print(f"  → Combined comparison image: {out_path}")

# Define a more general font type hint
AnyFont = Union[ImageFont.FreeTypeFont, ImageFont.ImageFont]

def get_scalable_font(font_size: int) -> AnyFont:
    """
    Try to return a scalable TTF font.
    It searches in order through Pillow's built-in directory, common Linux and macOS/Windows paths.
    If none are found, fall back to Pillow's default font (and try to set size).
    """
    # 1) Pillow built-in fonts
    try:
        # os.path.dirname(PIL.__file__) may be inaccurate in some environments
        # Pillow >= 9.3.0 provides a more reliable way
        from PIL import features
        if features.check("raqm"):
             pil_fonts_path = os.path.join(os.path.dirname(features.get_supported_modules()["raqm"].__file__), "DejaVuSans.ttf")
             if os.path.exists(pil_fonts_path):
                 return ImageFont.truetype(pil_fonts_path, font_size)

        # Traditional method
        pil_fonts_dir = os.path.join(os.path.dirname(PIL.__file__), "fonts")
        candidate = os.path.join(pil_fonts_dir, "DejaVuSans.ttf")
        if os.path.exists(candidate):
            return ImageFont.truetype(candidate, font_size)
    except Exception:
        pass

    # 2) Common Linux/macOS paths
    font_candidates = [
        # Linux
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
        # macOS
        "/System/Library/Fonts/Helvetica.ttc",
        "/Library/Fonts/Arial.ttf",
    ]
    # 3) Windows paths
    if os.name == 'nt':
        font_candidates.extend([
            "C:/Windows/Fonts/arial.ttf",
            "C:/Windows/Fonts/tahoma.ttf",
        ])

    for candidate in font_candidates:
        if os.path.exists(candidate):
            try:
                return ImageFont.truetype(candidate, font_size)
            except IOError:
                continue

    # 4) Final fallback
    print(f"[WARN] Cannot find any scalable TTF fonts, will use default font.")
    try:
        # Pillow >= 9.2.0 supports setting size for default font
        return ImageFont.load_default(size=font_size)
    except TypeError:
        # Compatible with older Pillow versions
        if font_size > 12: # Default font is small, give a hint if user wants large font
             print("[WARN] Current Pillow version is too old, cannot set default font size.")
        return ImageFont.load_default()


def annotate_image(img: Image.Image,
                   label: str,
                   top_margin: int = 8,
                   font_path: Optional[str] = None,
                   font_size: int = 36,
                   text_color=(0, 0, 0),
                   bg_color=(255, 255, 255)) -> Image.Image:
    """
    Draw label text (with background) at the top center of img, supports specifying font file and size.

    Args:
        img: PIL.Image.Image object to process.
        label: Text to draw.
        top_margin: Distance of text from top of image.
        font_path: (Optional) Path to .ttf or .otf font file. If provided, this font will be used preferentially.
        font_size: Text size.
        text_color: Text color.
        bg_color: Text background color.

    Returns:
        Processed PIL.Image.Image object.
    """
    draw = ImageDraw.Draw(img)
    font = None

    # 1) Load font: prioritize user-specified font_path
    if font_path and os.path.exists(font_path):
        try:
            font = ImageFont.truetype(font_path, font_size)
        except IOError:
            print(f"[WARN] Cannot load font from '{font_path}', will try automatic search.")

    # If user didn't specify font or loading failed, search automatically
    if font is None:
        font = get_scalable_font(font_size)

    # 2) Calculate text dimensions (compatible with different Pillow versions)
    try:
        # Pillow >= 10.0.0
        bbox = draw.textbbox((0, 0), label, font=font)
        text_w = bbox[2] - bbox[0]
        text_h = bbox[3] - bbox[1]
    except AttributeError:
        # Pillow < 10.0.0 (getsize is deprecated)
        text_w, text_h = draw.textsize(label, font=font)

    # 3) Calculate position: horizontally centered, top_margin from top
    img_w, img_h = img.size
    x = (img_w - text_w) // 2
    y = top_margin

    # 4) Draw background and text
    pad = max(4, font_size // 4)  # Padding dynamically adjusted based on font size
    draw.rectangle(
        [(x - pad, y - pad), (x + text_w + pad, y + text_h + pad)],
        fill=bg_color
    )
    draw.text((x, y - bbox[1]), label, fill=text_color, font=font) # Use bbox[1] to fine-tune vertical position
    
    return img

def make_pdf_from_images(root_dir: str,
                                      pdf_path: str,
                                      font_size: int = 48,
                                      title_margin: int = 20):
    """
    Stack the 4 compare images (front/side/top/iso) from each label_xxx directory under root_dir
    vertically into 4 rows, add a title row "Label ID: xxx" at the top,
    and finally compose a multi-page PDF (one page per label).
    """
    pages = []
    for label_folder in sorted(os.listdir(root_dir)):
        folder = os.path.join(root_dir, label_folder)
        if not os.path.isdir(folder) or not ('_label_' in label_folder):
            continue

        # Read PNG files from four viewpoints
        view_files = ['front_compare.png',
                      'side_compare.png',
                      'top_compare.png',
                      'iso_compare.png']
        imgs = []
        for vf in view_files:
            path = os.path.join(folder, vf)
            if not os.path.exists(path):
                raise FileNotFoundError(f"Missing {path}")
            imgs.append(Image.open(path).convert('RGB'))

        # Four images have the same dimensions
        w, h = imgs[0].size

        # Prepare font and title text
        font = get_scalable_font(font_size)
        # Use the full folder name as title (e.g., "gt_label_123_pred_456_789")
        title = label_folder

        # Calculate title dimensions
        dummy = ImageDraw.Draw(imgs[0])
        try:
            bbox = dummy.textbbox((0, 0), title, font=font)
            text_w = bbox[2] - bbox[0]
            text_h = bbox[3] - bbox[1]
        except AttributeError:
            text_w, text_h = font.getsize(title)

        # Title area height = text height + top and bottom margin
        title_h = text_h + title_margin * 2

        # Create full page canvas: width w, height = title_h + 4*h
        canvas = Image.new('RGB', (w, title_h + 4*h), (255, 255, 255))
        draw = ImageDraw.Draw(canvas)

        # Draw text in center of title area
        x = (w - text_w) // 2
        y = title_margin
        draw.text((x, y), title, fill=(0,0,0), font=font)

        # Paste four images row by row below title area
        for i, im in enumerate(imgs):
            canvas.paste(im, (0, title_h + i * h))

        pages.append(canvas)

    if not pages:
        print("[WARN] No label results found, PDF is empty.")
        return

    # Save multi-page PDF
    pages[0].save(
        pdf_path,
        save_all=True,
        append_images=pages[1:],
        resolution=300
    )
    print(f"[INFO] Generated PDF with title on each page: {pdf_path}")

def make_word_from_images(root_dir: str, word_path: str, error_type: str):
    """
    Save the 4 compare images (front/side/top/iso) from each label_xxx directory under root_dir
    to tables in a Word document, displaying two labels per page, with 4 views per label.
    """
    doc = Document()
    doc.add_heading(f'3D Visualization Results - {error_type.upper()}', 0)
    
    label_folders = sorted([f for f in os.listdir(root_dir) 
                           if os.path.isdir(os.path.join(root_dir, f)) and ('_label_' in f)])
    
    if not label_folders:
        print("[WARN] No label results found, Word document is empty.")
        return
    
    # Display two labels per page
    for page_idx in range(0, len(label_folders), 2):
        # Get labels for current page (maximum 2)
        current_labels = label_folders[page_idx:page_idx + 2]
        
        # Create a table for each label
        for label_idx, label_folder in enumerate(current_labels):
            folder = os.path.join(root_dir, label_folder)
            
            # Add Label title using the full folder name
            heading = doc.add_heading(f'{label_folder} ({error_type.upper()})', level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Create 2x2 table to display 4 views
            table = doc.add_table(rows=2, cols=2)
            table.style = 'Table Grid'
            
            # Define four viewpoints and their positions in the table
            views_positions = [
                ('front_compare.png', 'Front View', 0, 0),
                ('side_compare.png', 'Side View', 0, 1),
                ('top_compare.png', 'Top View', 1, 0),
                ('iso_compare.png', 'Isometric View', 1, 1)
            ]
            
            for img_file, view_name, row, col in views_positions:
                img_path = os.path.join(folder, img_file)
                cell = table.cell(row, col)
                
                if os.path.exists(img_path):
                    # Clear cell
                    cell.text = ''
                    
                    # Add viewpoint title
                    title_paragraph = cell.paragraphs[0]
                    title_paragraph.text = view_name
                    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    title_paragraph.runs[0].bold = True
                    
                    # Add image
                    img_paragraph = cell.add_paragraph()
                    img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    try:
                        # Resize image to fit table cell
                        img_paragraph.add_run().add_picture(img_path, width=Inches(3.0))
                    except Exception as e:
                        print(f"[ERROR] Cannot add image {img_path}: {e}")
                        img_paragraph.text = f"Image loading failed: {view_name}"
                else:
                    # If image doesn't exist, display error message
                    cell.text = f"{view_name}\nMissing image"
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add some spacing between two labels, but not at page end
            if label_idx == 0 and len(current_labels) > 1:
                doc.add_paragraph()  # Add empty paragraph as spacing
        
        # Add page break (except for last page)
        if page_idx + 2 < len(label_folders):
            doc.add_page_break()
    
    # Save Word document
    doc.save(word_path)
    print(f"[INFO] Generated Word document: {word_path}")

def main():
    p = argparse.ArgumentParser(
        description='Render gt and pred TIFF files with the same label_id from the same viewpoint and combine into a comparison image.'
    )
    p.add_argument('--input_dir', '-i', required=True, type=str, dest='input_dir',
                   help='Input directory containing FP and FN subfolders')
    p.add_argument('-s','--smooth', type=int, default=30,
                   help='Number of smoothing iterations (default 30; set 0 to skip smoothing)')
    p.add_argument('--spacing', nargs=3, type=float, default=[1.0, 1.0, 1.0],
                   help='Voxel spacing in z, y, x order (default: 1.0 1.0 1.0)')
    args = p.parse_args()

    spacing = tuple(args.spacing)
    print(f"[INFO] Using voxel spacing: {spacing} (z, y, x)")
    
    # Process two error types
    error_types = [('FN', 'fn'), ('FP', 'fp')]
    
    for folder_name, error_type in error_types:
        print(f"\n[INFO] Processing {error_type.upper()} error type...")
        
        # Build subfolder path
        error_input_dir = os.path.join(args.input_dir, folder_name)
        if not os.path.exists(error_input_dir):
            print(f"[WARN] {folder_name} folder not found, skipping.")
            continue
            
        # Match corresponding file pattern
        pattern = re.compile(rf'{error_type}_(\d+)_gt\.tiff$', re.IGNORECASE)
        
        # Create output directory for each error type under input directory
        error_output_dir = os.path.join(args.input_dir, f'3d_vis_{error_type}')
        os.makedirs(error_output_dir, exist_ok=True)

        for fname in sorted(os.listdir(error_input_dir)):
            m = pattern.match(fname)
            if not m:
                continue
            label_id = m.group(1)
            base = f'{error_type}_{label_id}_'
            print(f"  → {error_type.upper()} label_id: {label_id}")
            paths = {
                'gt':   os.path.join(error_input_dir, base + 'gt.tiff'),
                'pred': os.path.join(error_input_dir, base + 'pred.tiff'),
            }
            # Ensure both files exist
            if not all(os.path.exists(p) for p in paths.values()):
                print(f"[WARN] {error_type.upper()} label {label_id} missing gt/pred files, skipped.")
                continue

            # Read and check 3D data
            data_dict = {}
            for key, path in paths.items():
                try:
                    data = tifffile.imread(path)
                    if error_type == 'fp' and key == 'pred':
                        # If fp, filter out non-fp labels in pred
                        data[data != int(label_id)] = 0

                    if error_type == 'fn' and key == 'gt':
                        # If fn, filter out non-fn labels in gt
                        data[data != int(label_id)] = 0
                except Exception as e:
                    print(f"[ERROR] Failed to read {path}: {e}")
                    continue
                if data.ndim != 3:
                    print(f"[WARN] {path} is not 3D data, skipped.")
                    continue
                if not np.issubdtype(data.dtype, np.integer):
                    data = data.astype(np.uint8)
                data_dict[key] = data

            # Initialize top2_labels for directory naming
            top2_labels = []
            
            if error_type == 'fp':
                 # Extract gt label list where pred≠0
                 mask_pred0 = data_dict['pred'] == 0
                 gt_labels = np.unique(data_dict['gt'][~mask_pred0])
                 gt_labels = gt_labels[gt_labels != 0]  # Exclude background labels
                 
                 if len(gt_labels) > 0:
                     # Calculate overlap size between each label and pred foreground region
                     label_overlap_sizes = []
                     for label in gt_labels:
                         # Calculate overlap pixels between this label and pred foreground region
                         label_mask = data_dict['gt'] == label
                         pred_foreground = ~mask_pred0  # pred foreground region
                         overlap_size = np.sum(label_mask & pred_foreground)
                         label_overlap_sizes.append((label, overlap_size))
                     
                     # Sort by overlap size in descending order, take top 2
                     label_overlap_sizes.sort(key=lambda x: x[1], reverse=True)
                     top2_labels = [item[0] for item in label_overlap_sizes[:2]]
                     
                     # Keep only the top 2 labels with largest overlap
                     mask_gt_keep = np.isin(data_dict['gt'], top2_labels)
                     data_dict['gt'][~mask_gt_keep] = 0
                 else:
                     # If no overlapping labels, set all gt to zero
                     data_dict['gt'][:] = 0
                     
            elif error_type == 'fn':
                 # Extract pred label list where gt≠0
                 pred_labels = np.unique(data_dict['pred'][data_dict['gt'] != 0])
                 pred_labels = pred_labels[pred_labels != 0]  # Exclude background labels
                 
                 if len(pred_labels) > 0:
                     # Calculate overlap size between each pred label and gt
                     label_overlap_sizes = []
                     for label in pred_labels:
                         # Calculate overlap pixels between this pred label and gt foreground region
                         pred_label_mask = data_dict['pred'] == label
                         gt_foreground = data_dict['gt'] != 0  # gt foreground region
                         overlap_size = np.sum(pred_label_mask & gt_foreground)
                         label_overlap_sizes.append((label, overlap_size))
                     
                     # Sort by overlap size in descending order, take top 2
                     label_overlap_sizes.sort(key=lambda x: x[1], reverse=True)
                     top2_labels = [item[0] for item in label_overlap_sizes[:2]]
                     
                     # Keep only the top 2 pred labels with largest overlap
                     mask_pred_keep = np.isin(data_dict['pred'], top2_labels)
                     data_dict['pred'][~mask_pred_keep] = 0
                 else:
                     # If no overlapping labels, set all pred to zero
                     data_dict['pred'][:] = 0
                     
            if len(data_dict) != 2:
                continue

            # Create directory name based on error type and top2 labels
            if error_type == 'fn':
                # FN: gt_label_{label_id}_pred_{top2_labels[0]}_{top2_labels[1]}
                if len(top2_labels) >= 2:
                    dir_name = f'gt_label_{label_id}_pred_{top2_labels[0]}_{top2_labels[1]}'
                elif len(top2_labels) == 1:
                    dir_name = f'gt_label_{label_id}_pred_{top2_labels[0]}_none'
                else:
                    dir_name = f'gt_label_{label_id}_pred_none_none'
            else:  # error_type == 'fp'
                # FP: pred_label_{label_id}_gt_{top2_labels[0]}_{top2_labels[1]}
                if len(top2_labels) >= 2:
                    dir_name = f'pred_label_{label_id}_gt_{top2_labels[0]}_{top2_labels[1]}'
                elif len(top2_labels) == 1:
                    dir_name = f'pred_label_{label_id}_gt_{top2_labels[0]}_none'
                else:
                    dir_name = f'pred_label_{label_id}_gt_none_none'
             
            out_id_dir = os.path.join(error_output_dir, dir_name)
            os.makedirs(out_id_dir, exist_ok=True)
            print(f"[INFO] Processing {error_type.upper()} label {label_id} ...")

            # Four viewpoints
            views = {
                'front': (0, 0),
                'side':  (90, 0),
                'top':   (0, 90),
                'iso':   (45, 45),
            }

            for view_name, (az, el) in views.items():
                imgs = []
                # Corresponding text labels
                tags = ['GT', 'Pred']
                for key, tag in zip(('gt', 'pred'), tags):
                    img = render_single_view(data_dict[key], args.smooth, az, el, spacing=spacing)
                    # Add text to image
                    img = annotate_image(img, tag)
                    imgs.append(img)
                out_png = os.path.join(out_id_dir,
                                       f'{view_name}_compare.png')
                combine_modalities(imgs, out_png)
        
        # Generate PDF and Word documents for each error type
        pdf_file = os.path.join(error_output_dir, f'{error_type}_errors_all.pdf')
        word_file = os.path.join(error_output_dir, f'{error_type}_errors_all.docx')
        
        make_pdf_from_images(error_output_dir, pdf_file,
                            font_size=48,     # Adjustable title size
                            title_margin=20)
        make_word_from_images(error_output_dir, word_file, error_type)
if __name__ == '__main__':
    main()